from docx import Document
from docx2pdf import convert
import os

# === CONFIG ===
input_file = "kbr_coc_certs.docx"   # change this to your actual file name
output_folder = "Certificates"

# Create output folder
os.makedirs(output_folder, exist_ok=True)

# Load the big document
doc = Document(input_file)

current_doc = Document()
section_num = 1

for para in doc.paragraphs:
    # Add paragraph to current doc
    current_doc.add_paragraph(para.text)
    
    # Check for a manual page break in the paragraph XML
    if 'w:br w:type="page"' in para._element.xml:
        # Save current docx
        docx_filename = os.path.join(output_folder, f"certificate_{section_num}.docx")
        current_doc.save(docx_filename)
        
        # Convert to PDF
        convert(docx_filename)
        
        print(f"✅ Saved: certificate_{section_num}.docx and PDF")
        
        # Start a new doc
        current_doc = Document()
        section_num += 1

# Save last one if it has any content
if len(current_doc.paragraphs) > 0:
    docx_filename = os.path.join(output_folder, f"certificate_{section_num}.docx")
    current_doc.save(docx_filename)
    convert(docx_filename)
    print(f"✅ Saved: certificate_{section_num}.docx and PDF")

print("🎉 All certificates split and converted to PDF!")