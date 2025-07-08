from docx import Document
import os

# Load original document
doc = Document("kbr_coc_certs.docx")

output_folder = "Certificates"
os.makedirs(output_folder, exist_ok=True)

# Start new certificate document
current_doc = Document()
section_num = 1

for para in doc.paragraphs:
    current_doc.add_paragraph(para.text)
    
    # Check for a manual page break (common alternative to section break)
    if 'PageBreak' in para._element.xml:
        # Save current document
        filename = os.path.join(output_folder, f"certificate_{section_num}.docx")
        current_doc.save(filename)
        section_num += 1
        # Start a new document
        current_doc = Document()

# Save the last one
if len(current_doc.paragraphs) > 0:
    filename = os.path.join(output_folder, f"certificate_{section_num}.docx")
    current_doc.save(filename)

print("✅ All certificates saved as separate .docx files.")